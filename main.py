import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
import webbrowser

def obtener_todos_los_parrafos(document):
    
    parrafos = []

    parrafos.extend(document.paragraphs)

    for i, para in enumerate(parrafos):
        print(f"Indice {i}: {para.text}")
        print(para._element.xml)  # Esto muestra el XML interno
        print("-" * 40)


    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                
                parrafos.extend(cell.paragraphs)
            
    
    for section in document.sections:
        parrafos.extend(section.header.paragraphs)
        parrafos.extend(section.footer.paragraphs)

    return parrafos

# Función para extraer los links (texto visible y URL) de un archivo .docx
def extraer_links_docx(file_path):
    document = Document(file_path)
    links = []

    # Recorremos cada parrafo
    for para in obtener_todos_los_parrafos(document):
        para_xml = para._element
        

        # Caso 1
        hyperlink_elements = para_xml.xpath(".//w:hyperlink") 
        for hyperlink in hyperlink_elements:
            
            r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id") # Aca obtendriamos el elemento XML, que tiene atributos como "target", en "target" se almacena el URL
            

            if r_id:
                try:
                    url = document.part.rels[r_id].target_ref # En esta parte obtendriamos ya el URL

                    text_parts = []
                    # En este fragmento obtendriamos el nombre del elemento hyperlink XML
                    for node in hyperlink.xpath(".//w:t"):
                        if node.text: # Nueva linea
                            text_parts.append(node.text)
                    texto_visible = ''.join(text_parts)
                    #print(texto_visible)
                    if texto_visible:
                        links.append((texto_visible, url))
                except KeyError:
                    continue
        
        # Caso 2
        instr_text_nodes = para_xml.xpath('.//w:instrText')
        for node in instr_text_nodes:
            if node.text and 'HYPERLINK' in node.text:
                # Extrae la URL entre comillas
                partes = node.text.split('"')
                if len(partes) >= 2:
                    url = partes[1]
                    # Obtener texto visible del párrafo
                    text_parts = [t.text for t in para_xml.xpath('.//w:t') if t.text]
                    texto_visible = ''.join(text_parts)
                    # Verificamos que no esté duplicado
                    if texto_visible:
                        links.append((texto_visible, url))        

    return links

# Mostrar los links en la tabla de la interfaz
def mostrar_links_en_tabla(links):
    for row in tabla.get_children():
        tabla.delete(row)

    for texto_visible, url in links:
        tabla.insert("", "end", values=(texto_visible, url))

# Función para abrir el enlace en el navegador
def abrir_link(event):
    item = tabla.selection()  # Obtener la fila seleccionada
    if item:
        url = tabla.item(item, "values")[1]  # Obtener el link
        if url:
            webbrowser.open(url)  # Abrir en el navegador

# Selección del archivo y procesamiento
def open_file_dialog():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        try:
            links = extraer_links_docx(file_path)

            if not links:
                messagebox.showinfo("Sin enlaces", "No se encontraron enlaces en el documento.")
            else:
                mostrar_links_en_tabla(links)

        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error al procesar el archivo:\n{str(e)}")

# Interfaz con Tkinter
root = tk.Tk()
root.title("Extractor de links")
root.geometry("700x500")
root.configure(bg="#ecf0f1")

titulo = tk.Label(root, text="Catter", font=("Arial", 18, "bold"), bg="#ecf0f1")
titulo.pack(pady=10)

boton_abrir = tk.Button(root, text="Seleccionar archivo", command=open_file_dialog, bg="#3498db", fg="white", font=("Arial", 12, "bold"))
boton_abrir.pack(pady=10)

# Tabla para mostrar los enlaces
tabla = ttk.Treeview(root, columns=("Texto", "Link"), show="headings")
tabla.heading("Texto", text="Tag")
tabla.heading("Link", text="Enlace")
tabla.column("Texto", width=200, anchor="w")
tabla.column("Link", width=450, anchor="w")
tabla.pack(pady=20, fill=tk.BOTH, expand=True)

# Hacer que el clic en la fila abra el enlace en el navegador
tabla.bind("<Double-1>", abrir_link)

root.mainloop()
